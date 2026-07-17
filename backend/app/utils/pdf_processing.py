import io
from pathlib import Path
from typing import Optional
import fitz
from PIL import Image
import pypdf
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def pdf_to_images(pdf_path: Path, dpi: int = 200) -> list[Image.Image]:
    doc = fitz.open(pdf_path)
    images = []
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    for page in doc:
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    doc.close()
    return images


def pdf_to_images_bytes(pdf_path: Path, fmt: str = "jpeg", dpi: int = 200) -> list[bytes]:
    images = pdf_to_images(pdf_path, dpi)
    bufs = []
    for img in images:
        buf = io.BytesIO()
        img.save(buf, format=fmt.upper(), quality=92)
        bufs.append(buf.getvalue())
    return bufs


def pdf_to_text(pdf_path: Path) -> str:
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def pdf_to_docx(pdf_path: Path) -> bytes:
    doc = fitz.open(pdf_path)
    output = Document()
    for page_num, page in enumerate(doc):
        blocks = page.get_text("blocks")
        blocks.sort(key=lambda b: (b[1], b[0]))
        for block in blocks:
            text = block[4].strip()
            if not text:
                continue
            if page_num > 0 or blocks.index(block) > 0:
                output.add_paragraph(text)
            else:
                output.add_paragraph(text)
    doc.close()
    buf = io.BytesIO()
    output.save(buf)
    return buf.getvalue()


def compress_pdf(input_path: Path, quality: int = 0) -> bytes:
    reader = pypdf.PdfReader(input_path)
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    for page in writer.pages:
        for img in page.images:
            img.replace(img.image, quality=quality)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def extract_images_from_pdf(pdf_path: Path, min_size: int = 100) -> list[tuple[bytes, str]]:
    doc = fitz.open(pdf_path)
    results = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            ext = base_image["ext"]
            img_pil = Image.open(io.BytesIO(img_bytes))
            if img_pil.width >= min_size and img_pil.height >= min_size:
                results.append((img_bytes, ext))
    doc.close()
    return results
